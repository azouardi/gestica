
from decimal import Decimal
from audits.models import Conclusion, Lead, Section
from ordres.models import Ordre
import ordres
import docx
from docx.opc.oxml import qn
from docx.oxml import OxmlElement
from customers.models import Company
from editions.models import DocModel, DocModelWord
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, render

from django.http import HttpResponse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
# Create your views here.
from dateutil.relativedelta import relativedelta

from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE

from tasks.models import Situation

# import locale
# locale.setlocale(locale.LC_NUMERIC, 'de_DE.utf-8')
# locale.setlocale(locale.LC_ALL, 'de_DE.utf-8')


def writingword(request, pk, code_docmodel):
    company=Company.objects.get(pk=pk)
    model = DocModel.objects.get(code_docmodel=code_docmodel)
    modelitem = DocModelWord.objects.filter(code_docmodel=code_docmodel).order_by('ordre_parag')
    document = Document()
    
    alignment_dict = {'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  'left': WD_ALIGN_PARAGRAPH.LEFT}

    orientation_dict = {'portrait': WD_ORIENTATION.PORTRAIT,
                    'landscape': WD_ORIENTATION.LANDSCAPE}
  

    def change_orientation(orientation='portrait', set_left_margin=1.0, set_right_margin=1.0):
        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = orientation_dict[orientation]
        section.page_width = new_width
        section.page_height = new_height
        section.left_margin = Inches(set_left_margin)
        section.right_margin = Inches(set_right_margin)


    def add_logo(path, align):
        document.add_picture('media/'+str(company.file), width=Inches(4.5), height=Inches(1.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = alignment_dict[align]


    def add_content(content, space_after=6, font_name='Arial', font_size=11, line_spacing=1.25, space_before=3,
                align='justify', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        paragraph = document.add_paragraph(content)
        run = paragraph.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def add_content_cell_bis(htable, row, cell, content, space_after=2, font_name='Arial', font_size=9, line_spacing=1, space_before=0,
                align='left', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        
        htab_cells=htable.rows[row].cells
        ht0=htab_cells[cell]
        ht0.text = content
        paragraph = ht0.paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def add_run(content, font_name='Arial', set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        paragraph = document.paragraphs[-1]
        paragraph.add_run(content)
        run = paragraph.runs[-1]
        font = run.font
        font.name = font_name
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline


    def create_numbered_list():
        pass
    

    def add_subheading(subheading, level):
        document.add_heading(subheading, level)

    # def entetepage():
    #     header = document.sections[0].header
    #     htable=header.add_table(1, 3, Inches(10))
    #     htab_cells=htable.rows[0].cells
    #     ht0=htab_cells[0].add_paragraph()
    #     kh=ht0.add_run()
    #     kh.add_picture('media/'+str(company.file), width=Inches(2))
    #     hp1=htab_cells[2].add_paragraph(str(company.name))
    #     hp1.paragraph_format.space_before = Pt(0)
    #     hp1.paragraph_format.space_after = Pt(0)
    #     hp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #     hp1=htab_cells[2].add_paragraph(str(company.adresse))
    #     hp1.paragraph_format.space_before = Pt(0)
    #     hp1.paragraph_format.space_after = Pt(0)
    #     hp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #     hp1=htab_cells[2].add_paragraph(str(company.city))
    #     hp1.paragraph_format.space_before = Pt(0)
    #     hp1.paragraph_format.space_after = Pt(0)
    #     hp1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def entetepage():
        header = document.sections[0].header
        htable=header.add_table(5, 2, Inches(12))
        a = htable.cell(0,0)
        b = htable.cell(4,0)
        A = a.merge(b)
        ht0=A.add_paragraph()
        kh=ht0.add_run()
        kh.add_picture('media/'+str(company.file), width=Inches(2))
        for cell in htable.columns[0].cells:
            cell.width = Inches(8)
        add_content_cell_bis(htable=htable, row=1,cell=1, content=str(company.name), space_before=10, space_after=6, set_bold=True)
        add_content_cell_bis(htable=htable, row=2,cell=1, content=str(company.adresse), space_after=6, set_bold=True)
        add_content_cell_bis(htable=htable, row=3,cell=1, content=str(company.city), space_after=6, set_bold=True)


        
    def piedpage():
        footer = document.sections[0].footer
        fp1 =footer.paragraphs[-1]
        fp1.add_run("_____________________________________________________________")
        fp1.paragraph_format.space_before = Pt(0)
        fp1.paragraph_format.space_after = Pt(3)
        fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fp1 = footer.add_paragraph(str(company.name))
        fp1.paragraph_format.space_before = Pt(0)
        fp1.paragraph_format.space_after = Pt(0)
        fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fp1 = footer.add_paragraph(str(company.adresse) +' - '+ str(company.city))
        fp1.paragraph_format.space_before = Pt(0)
        fp1.paragraph_format.space_after = Pt(0)
        fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fp1 = footer.add_paragraph('RC: ' + str(company.rc) +' IF: '+ str(company.identifiant_fiscal) +' TP: '+ str(company.tp) +' ICE: '+ str(company.ice))
        fp1.paragraph_format.space_before = Pt(0)
        fp1.paragraph_format.space_after = Pt(0)
        fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# début de la rédaction


    def modelword(code_docmodel):

        sections = document.sections
        for section in sections:
            section.header_distance = Cm(0.5)
            section.footer_distance = Cm(0.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            
        if model.header == True:
            entetepage()

        if model.footer == True:
            piedpage()
        
        for i in modelitem:
            type_parag = i.type_parag
            texte = i.texte
            if type_parag == 1:
                add_subheading(str(texte), level=i.sort_parag)
            if type_parag == 2:
                add_content(str(texte), align=i.alignment_parag, space_after=i.space_after_parag, space_before=i.space_befor_parag, font_size=i.size_font, set_bold=i.bold_font, set_italic=i.italic_font, set_underline=i.undeline_font)            
            if type_parag == 3:
                add_run(str(texte), set_bold=i.bold_font, set_italic=i.italic_font, set_underline=i.undeline_font) 
                           
    modelword(code_docmodel)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename='+str(company.name)+'_'+str(model)+'_Lettre.docx'
    document.save(response)
    return response


def modelCommun(request, pk, code_docmodel):
    redaction = writingword(request=request, pk=pk, code_docmodel=code_docmodel)
    return redaction

def writingdevis(request, pk):
    ordre=Ordre.objects.get(pk=pk)
    serviceitems = ordre.serviceitem_set.all()
    outlayitems = ordre.outlayitem_set.all()
    document = Document()

    alignment_dict = {'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  'left': WD_PARAGRAPH_ALIGNMENT.LEFT}

    orientation_dict = {'portrait': WD_ORIENTATION.PORTRAIT,
                    'landscape': WD_ORIENTATION.LANDSCAPE}
  
    def add_logo(path, align):
        document.add_picture(path, width=Inches(4.5), height=Inches(1.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = alignment_dict[align]
        
    def add_content(content, space_after=6, font_name='Arial', font_size=11, line_spacing=1.25, space_before=3,
                align='justify', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        paragraph = document.add_paragraph(content)
        run = paragraph.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def add_content_cell(htab_cells, content, space_after=2, font_name='Arial', font_size=9, line_spacing=1, space_before=0,
                align='left', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        hp=htab_cells.add_paragraph(content)
        run = hp.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = hp.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control


    def add_content_cell_bis(htable, row, cell, content, space_after=2, font_name='Arial', font_size=9, line_spacing=1, space_before=0,
                align='left', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        
        htab_cells=htable.rows[row].cells
        ht0=htab_cells[cell]
        ht0.text = content
        paragraph = ht0.paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def add_run(content, font_name='Arial', set_bold=True, set_italic=False, set_underline=False, set_all_caps=False):
        paragraph = document.paragraphs[-1]
        paragraph.add_run(content)
        run = paragraph.runs[-1]
        font = run.font
        font.name = font_name
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline


    def add_subheading(subheading, level):
        document.add_heading(subheading, level)

    def entetepage():
        header = document.sections[0].header
        htable=header.add_table(5, 2, Inches(12))
        a = htable.cell(0,0)
        b = htable.cell(4,0)
        A = a.merge(b)
        ht0=A.add_paragraph()
        kh=ht0.add_run()
        kh.add_picture('media/'+str(ordre.office.company.file), width=Inches(3))
        for cell in htable.columns[0].cells:
            cell.width = Inches(8)
        add_content_cell_bis(htable=htable, row=0,cell=1, content='EXPERTISE COMPTABLE, AUDIT', space_before=10, space_after=6, set_bold=True)
        add_content_cell_bis(htable=htable, row=1,cell=1, content='COMMISSARIAT AUX COMPTES', space_after=6, set_bold=True)
        add_content_cell_bis(htable=htable, row=2,cell=1, content='CONSEIL JURIDIQUE ET FISCAL', space_after=6, set_bold=True)
        add_content_cell_bis(htable=htable, row=3,cell=1, content='DROIT DU TRAVAIL ET PAIE', space_after=6, set_bold=True)
        add_content_cell_bis(htable=htable, row=4,cell=1, content='TRANSMISSION D’ENTREPRISES', space_after=6, set_bold=True)



        
    def piedpage():
        footer = document.sections[0].footer
        fp1 =footer.paragraphs[-1]
        fp1.add_run("______________________________________________________________________________________________")
        run =fp1.runs
        run[0].font.size = Pt(14)
        run[0].font.bold = True
        run[0].font.color.rgb = RGBColor(255, 102, 0)
        ftable=footer.add_table(4, 3, Inches(15))
        a = ftable.cell(0,0)
        b = ftable.cell(1,0)
        A = a.merge(b)
        for cell in ftable.columns[0].cells:
            cell.width = Inches(6)  
        for cell in ftable.columns[1].cells:
            cell.width = Inches(8) 
        add_content_cell_bis(htable=ftable, row=0,cell=0,content=str(ordre.office.company.adresse)+ ' - '+str(ordre.office.company.city))
        add_content_cell_bis(htable=ftable, row=0,cell=1,content=str(ordre.office.company.name))
        add_content_cell_bis(htable=ftable, row=0,cell=2,content='ouardi@eurodefis.com')
        
        add_content_cell_bis(htable=ftable, row=1,cell=1,content='RC: '+str(ordre.office.company.rc)+ ' - CNSS:'+str(ordre.office.company.cnss))
        add_content_cell_bis(htable=ftable, row=1,cell=2,content='arji@eurodefis.com')
        
        add_content_cell_bis(htable=ftable, row=2,cell=0,content='Tél  : +212 522 58 48 69')
        add_content_cell_bis(htable=ftable, row=2,cell=1,content='TP: '+str(ordre.office.company.tp)+ ' - IF:'+str(ordre.office.company.identifiant_fiscal))
        add_content_cell_bis(htable=ftable, row=2,cell=2,content=' ')

        add_content_cell_bis(htable=ftable, row=3,cell=0,content='Fax : +212 522 33 51 70')
        add_content_cell_bis(htable=ftable, row=3,cell=1,content='ICE: '+str(ordre.office.company.ice))
        add_content_cell_bis(htable=ftable, row=3,cell=2,content='www.eurodefis.com')        

# début de la rédaction


    def modelword(request):

        sections = document.sections
        for section in sections:
            section.header_distance = Cm(0.5)
            section.footer_distance = Cm(0.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            
        entetepage()

        piedpage()
        tableT = document.add_table(rows=1, cols=2)
        tab_cells=tableT.rows[0].cells
      
        add_content_cell(htab_cells=tab_cells[1],content=str(ordre.company.name), set_bold=True)
        add_content_cell(htab_cells=tab_cells[1],content=str(ordre.company.adresse) +' - '+str(ordre.company.city), font_size=9)
        add_content_cell(htab_cells=tab_cells[1],content=str(ordre.company.ice), set_bold=True)
        
        add_content(content='Devis d\'honoraire', font_size=15, set_bold=True, align='center',set_underline=True, space_after=12, space_before=12) 
        
        #Tableau honoraires et débours
        tableD = document.add_table(rows=len(serviceitems)+2, cols=5, style='Table Grid')
        for cell in tableD.columns[1].cells:
            cell.width = Inches(6)
        for cell in tableD.columns[2].cells:
            cell.width = Inches(1) 
        tab_cells=tableD.rows[0].cells

        add_content_cell(htab_cells=tab_cells[0],content='Référence', font_size=9, set_all_caps=False)
        add_content_cell(htab_cells=tab_cells[1],content='Désignation', font_size=9, set_all_caps=False)
        add_content_cell(htab_cells=tab_cells[2],content='Quantité', font_size=9,  set_all_caps=False)
        add_content_cell(htab_cells=tab_cells[3],content='Prix Unit.', font_size=9, set_all_caps=False)
        add_content_cell(htab_cells=tab_cells[4],content='Montant HT', font_size=9, set_all_caps=False)      
        
        #Honoriares
        r = 1
        sumHon = 0
        sumHonTVA = 0
        for item in serviceitems: 
                  
            tab_cells=tableD.rows[r].cells      
            add_content_cell(htab_cells=tab_cells[0],content=str(item.service.moduleservice), font_size=9, set_bold=False, set_all_caps=False)
            if item == item:
                pass
            else:
                add_content_cell(htab_cells=tab_cells[0],content=str(item.service.moduleservice), font_size=9, set_bold=False, set_all_caps=False)
            add_content_cell(htab_cells=tab_cells[1],content=str(item.service.text), font_size=9, set_bold=False, set_all_caps=False)
            for desc in item.service.servicedescription_set.all():
                add_content_cell(htab_cells=tab_cells[1],content=str(desc.text), font_size=8, set_bold=False, set_all_caps=False)            
            add_content_cell(htab_cells=tab_cells[2],content=str(item.quantity), font_size=9, align='right', set_bold=False, set_all_caps=False)            
            add_content_cell(htab_cells=tab_cells[3],content=str('{:8,.2f}'.format(item.amount)), font_size=9, align='right', set_bold=False, set_all_caps=False)            
            add_content_cell(htab_cells=tab_cells[4],content=str('{:8,.2f}'.format(item.amount* item.quantity)), font_size=9, align='right', set_bold=False, set_all_caps=False)            
            r = r + 1
            sumHon = sumHon + (item.amount* item.quantity)
            sumHonTVA = sumHonTVA + (item.amount* item.quantity * item.taxe)
    
        #Débours
        r = len(serviceitems)+1
        sumDeb = 0
        tab_cells=tableD.rows[r].cells      
        add_content_cell(htab_cells=tab_cells[0],content='DÉBOURS', font_size=9, set_bold=False, set_all_caps=False)
        for item in outlayitems: 
            add_content_cell(htab_cells=tab_cells[1],content=str(item.outlay.text), font_size=9, set_bold=False, set_all_caps=False)         
            add_content_cell(htab_cells=tab_cells[2],content=str(item.quantity), font_size=9, align='right', set_bold=False, set_all_caps=False)            
            add_content_cell(htab_cells=tab_cells[3],content=str('{:8,.2f}'.format(item.amount)), font_size=9, align='right', set_bold=False, set_all_caps=False)            
            add_content_cell(htab_cells=tab_cells[4],content=str('{:8,.2f}'.format(item.amount* item.quantity)), font_size=9, align='right', set_bold=False, set_all_caps=False)            
            # r = r + 1
            sumDeb = sumDeb + (item.amount* item.quantity)
    
    
            
        add_content(content=' ', font_size=15, set_bold=True, align='left', space_after=0, space_before=0)
        
        tableS = document.add_table(rows=4, cols=5, style='Table Grid')
        for cell in tableS.columns[0].cells:
            cell.width = Inches(6)
        for cell in tableS.columns[2].cells:
            cell.width = Inches(1)    
        add_content_cell_bis(htable=tableS, row=0,cell=0, content=' ')
        add_content_cell_bis(htable=tableS, row=0,cell=1, content='Montant HT', align='right')
        add_content_cell_bis(htable=tableS, row=0,cell=2,content='Taux', align='right')
        add_content_cell_bis(htable=tableS, row=0,cell=3,content='TVA', align='right')
        add_content_cell_bis(htable=tableS, row=0,cell=4,content='Total', align='right', set_bold=True)
        
        add_content_cell_bis(htable=tableS, row=1,cell=0, content='Honoraires')
        add_content_cell_bis(htable=tableS, row=1,cell=1,content=str('{:8,.2f}'.format(sumHon)), align='right')
        add_content_cell_bis(htable=tableS, row=1,cell=2, content='20%', align='right')
        add_content_cell_bis(htable=tableS, row=1,cell=3,content=str('{:8,.2f}'.format(sumHonTVA)), align='right')
        add_content_cell_bis(htable=tableS, row=1,cell=4,content=str('{:8,.2f}'.format(sumHon + sumHonTVA)), align='right', set_bold=True)

        add_content_cell_bis(htable=tableS, row=2,cell=0, content='Débours ')
        add_content_cell_bis(htable=tableS, row=2,cell=1,content=str('{:8,.2f}'.format(sumDeb)), align='right')
        add_content_cell_bis(htable=tableS, row=2,cell=2, content='---', align='right')
        add_content_cell_bis(htable=tableS, row=2,cell=3,content='---', align='right')
        add_content_cell_bis(htable=tableS, row=2,cell=4,content=str('{:8,.2f}'.format(sumDeb)), align='right', set_bold=True)    
        
        add_content_cell_bis(htable=tableS, row=3,cell=0, content='Totaux ', set_bold=True, align='right')
        add_content_cell_bis(htable=tableS, row=3,cell=1,content=str('{:8,.2f}'.format(sumDeb+sumHon)), align='right', set_bold=True)
        add_content_cell_bis(htable=tableS, row=3,cell=2, content='---', align='right')
        add_content_cell_bis(htable=tableS, row=3,cell=3,content=str('{:8,.2f}'.format(sumHonTVA)), align='right', set_bold=True)
        add_content_cell_bis(htable=tableS, row=3,cell=4,content=str('{:8,.2f}'.format(sumDeb + sumHon + sumHonTVA)), align='right', set_bold=True)    

        add_content(content='Merci de libeller votre chèque ou virement', font_size=9, set_italic=True, space_after=0,space_before=12)
        add_content(content='Bénéficiaire : '+ str(ordre.office.company.name), font_size=9, set_bold=True, space_after=0,space_before=0)
        banks= ordre.office.company.bankacount_set.all()
        for bank in banks:
            add_content(content='Banque : '+str(bank.bank_name), set_bold=False, font_size=9,space_after=0, space_before=0)
            add_content(content='Agence : '+str(bank.agency), set_bold=False, font_size=9,space_after=0, space_before=0)
            add_content(content='RIB : '+str(bank.rib), set_bold=True, font_size=9,space_after=0, space_before=0)

    
    modelword(request)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Devis.docx'
    document.save(response)
    return response

def modelDevis(request, pk):
    redaction = writingdevis(request=request, pk=pk)
    return redaction



#-----------------NOTE SYNTHESE-----------------
def writingnotesynthese(request, pk):
    s=Situation.objects.get(pk=pk)
    sections = s.section_set.all()
    document = Document()

    alignment_dict = {'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                  'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  'left': WD_PARAGRAPH_ALIGNMENT.LEFT}

    orientation_dict = {'portrait': WD_ORIENTATION.PORTRAIT,
                    'landscape': WD_ORIENTATION.LANDSCAPE}
  
    def add_logo(path, align):
        document.add_picture(path, width=Inches(4.5), height=Inches(1.5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = alignment_dict[align]
        
    def add_content(content, space_after=6, font_name='Arial', font_size=12, line_spacing=1.5, space_before=6,
                align='justify', keep_together=False, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        paragraph = document.add_paragraph(content)
        run = paragraph.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def add_content_cell(htab_cells, content, space_after=2, font_name='Arial', font_size=9, line_spacing=1, space_before=0,
                align='left', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        hp=htab_cells.add_paragraph(content)
        run = hp.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = hp.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control


    def add_content_cell_bis(htable, row, cell, content, space_after=2, font_name='Arial', font_size=9, line_spacing=1, space_before=0,
                align='left', keep_together=True, keep_with_next=False, page_break_before=False,
                widow_control=False, set_bold=False, set_italic=False, set_underline=False, set_all_caps=False):
        
        htab_cells=htable.rows[row].cells
        ht0=htab_cells[cell]
        ht0.text = content
        paragraph = ht0.paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = alignment_dict.get(align.lower())
        paragraph_format.space_before = Pt(space_before)
        paragraph_format.space_after = Pt(space_after)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control

    def add_run(content, font_name='Arial', set_bold=True, set_italic=False, set_underline=False, set_all_caps=False):
        paragraph = document.paragraphs[-1]
        paragraph.add_run(content)
        run = paragraph.runs[-1]
        font = run.font
        font.name = font_name
        font.bold = set_bold
        font.italic = set_italic
        font.all_caps = set_all_caps
        font.underline = set_underline


    def add_subheading(subheading, level):
        document.add_heading(subheading, level)

    def entetepage():
        header = document.sections[0].header
        htable=header.add_table(5, 2, Inches(12))
        a = htable.cell(0,0)
        b = htable.cell(4,0)
        A = a.merge(b)
        ht0=A.add_paragraph()
        kh=ht0.add_run()
        kh.add_picture('media/'+str(s.lettremission.office.company.file), width=Inches(3))
        for cell in htable.columns[0].cells:
            cell.width = Inches(8)
        add_content_cell_bis(htable=htable, row=0,cell=1, content='EXPERTISE COMPTABLE, AUDIT', space_before=0, space_after=3, set_bold=True)
        add_content_cell_bis(htable=htable, row=1,cell=1, content='COMMISSARIAT AUX COMPTES', space_after=3, set_bold=True)
        add_content_cell_bis(htable=htable, row=2,cell=1, content='CONSEIL JURIDIQUE ET FISCAL', space_after=3, set_bold=True)
        add_content_cell_bis(htable=htable, row=3,cell=1, content='DROIT DU TRAVAIL ET PAIE', space_after=3, set_bold=True)
        add_content_cell_bis(htable=htable, row=4,cell=1, content='TRANSMISSION D’ENTREPRISES', space_after=3, set_bold=True)



        
    def piedpage():
        footer = document.sections[0].footer
        fp1 =footer.paragraphs[-1]
        fp1.add_run("______________________________________________________________________________________________")
        run =fp1.runs
        run[0].font.size = Pt(14)
        run[0].font.bold = True
        run[0].font.color.rgb = RGBColor(255, 102, 0)
        ftable=footer.add_table(4, 3, Inches(15))
        a = ftable.cell(0,0)
        b = ftable.cell(1,0)
        A = a.merge(b)
        for cell in ftable.columns[0].cells:
            cell.width = Inches(6)  
        for cell in ftable.columns[1].cells:
            cell.width = Inches(8) 
        add_content_cell_bis(htable=ftable, row=0,cell=0,content=str(s.lettremission.office.company.adresse)+ ' - '+str(s.lettremission.office.company.city))
        add_content_cell_bis(htable=ftable, row=0,cell=1,content=str(s.lettremission.office.company.name))
        add_content_cell_bis(htable=ftable, row=0,cell=2,content='ouardi@eurodefis.com')
        
        add_content_cell_bis(htable=ftable, row=1,cell=1,content='RC: '+str(s.lettremission.office.company.rc)+ ' - CNSS:'+str(s.lettremission.office.company.cnss))
        add_content_cell_bis(htable=ftable, row=1,cell=2,content='arji@eurodefis.com')
        
        add_content_cell_bis(htable=ftable, row=2,cell=0,content='Tél  : +212 522 58 48 69')
        add_content_cell_bis(htable=ftable, row=2,cell=1,content='TP: '+str(s.lettremission.office.company.tp)+ ' - IF:'+str(s.lettremission.office.company.identifiant_fiscal))
        add_content_cell_bis(htable=ftable, row=2,cell=2,content=' ')

        add_content_cell_bis(htable=ftable, row=3,cell=0,content='Fax : +212 522 33 51 70')
        add_content_cell_bis(htable=ftable, row=3,cell=1,content='ICE: '+str(s.lettremission.office.company.ice))
        add_content_cell_bis(htable=ftable, row=3,cell=2,content='www.eurodefis.com')        

# début de la rédaction


    def modelword(request):

        sections = document.sections
        for section in sections:
            section.header_distance = Cm(0.5)
            section.footer_distance = Cm(0.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            
        entetepage()
        piedpage()
        
        #-------------page de garde--------------
        add_content(content='Note de synthèse', font_size=24, set_bold=True, align='center',set_underline=False, space_after=32, space_before=64, set_all_caps=True) 
        add_content(content='Au titre de la période cloturée au :'+str(s.date_closing), font_size=14, set_bold=True, align='center',set_underline=False, space_after=32, space_before=32) 
        add_content(content=str(s.lettremission.company.name), font_size=18, set_bold=True, align='center',set_underline=False, space_after=6, space_before=32, set_all_caps=True) 
        add_logo(path='media/'+str(s.lettremission.company.file),align='center')
        document.add_picture('media/'+str(s.lettremission.company.file), width=Inches(3))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = alignment_dict['center']

        # -------------saut de page--------------------
        document.add_paragraph()
        last_paragraph = document.paragraphs[-1]
        run = last_paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)
        
        #-----------------CONTENU-------------------
        
        sections= Section.objects.filter(situation=s).filter(statut__in=[1,2]).order_by('section')
        if sections:
            date_closing_n1=s.date_closing+relativedelta(years=-1)

                   #Tableau ESG-----------
            
            tableESG = document.add_table(rows=34, cols=6, style='Table Grid')
            for cell in tableESG.columns[0].cells:
                cell.width = Inches(5)
            for cell in tableESG.columns[2].cells:
                cell.width = Inches(2) 
            tab_cells=tableESG.rows[0].cells

            add_content_cell_bis(htable=tableESG, row=0, cell=0,content='Rubrique', font_size=12, set_all_caps=True, set_bold=True)
            add_content_cell_bis(htable=tableESG, row=0, cell=1,content='Note', font_size=12, set_all_caps=True, set_bold=True)
            add_content_cell_bis(htable=tableESG, row=0, cell=2,content=str(s.date_closing), font_size=12, set_all_caps=True, set_bold=True)
            add_content_cell_bis(htable=tableESG, row=0, cell=3,content=str(date_closing_n1), font_size=12, set_all_caps=True, set_bold=True)
            add_content_cell_bis(htable=tableESG, row=0, cell=4,content='Var.', font_size=12, set_all_caps=True, set_bold=True)      
            add_content_cell_bis(htable=tableESG, row=0, cell=5,content='%', font_size=12, set_all_caps=True, set_bold=True)      
            
            tab_cells=tableESG.rows[1].cells
            lead = Lead.objects.filter(code_account_id='711').filter(section__in=[sections[0].id])
            if lead:
                add_content_cell_bis(htable=tableESG, row=1, cell=0,content=str(lead[0].code_account.name_account), font_size=12, set_all_caps=False)
                add_content_cell_bis(htable=tableESG, row=1, cell=1,content='-', font_size=9, set_all_caps=False)
                add_content_cell_bis(htable=tableESG, row=1, cell=2,content=str('{:8,.2f}'.format(lead[0].data_n)), font_size=12,  set_all_caps=False)
                add_content_cell_bis(htable=tableESG, row=1, cell=3,content=str('{:8,.2f}'.format(lead[0].data_n1)), font_size=12, set_all_caps=False)
                add_content_cell_bis(htable=tableESG, row=1, cell=4,content=str('{:8,.2f}'.format(lead[0].data_delta)), font_size=12, set_all_caps=False)      
                add_content_cell_bis(htable=tableESG, row=1, cell=5,content=str('{:8,.2f}'.format(lead[0].data_delta_pc)), font_size=12, set_all_caps=False)               
            
            #-----------------REVUE ANALYTIQUE-------------------
            for sect in sections:
                revues = Conclusion.objects.filter(section=sect).filter(statut__in=[1,2])
                if revues:
                    add_subheading(str(sect.section.name_section),2)
                    for rev in revues:
                        add_content(content=str(rev.conclusion))

        
        
    modelword(request)
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=Devis.docx'
    document.save(response)
    return response

def modelNoteSynthese(request, pk):
    redaction = writingnotesynthese(request=request, pk=pk)
    return redaction